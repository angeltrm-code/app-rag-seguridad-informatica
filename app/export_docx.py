#!/usr/bin/env python3
"""
app/export_docx.py — Script fallback para exportar los planes MD a DOCX.
Utiliza python-docx para leer las secciones principales y tablas.

Uso:
    python -m app.export_docx
"""

import os
import sys
import re
from pathlib import Path

def sanitize_xml(text):
    if not isinstance(text, str): return text
    # Strip illegal XML characters
    return re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("[ERROR] python-docx no está instalado. Ejecute: pip install python-docx")
    sys.exit(1)

PROJECT_ROOT = Path(__file__).resolve().parent.parent

def extract_table(lines, start_idx):
    """Extrae una tabla markdown empezando en start_idx."""
    table_lines = []
    idx = start_idx
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith('|'):
            break
        table_lines.append(line)
        idx += 1
    return table_lines, idx

def parse_markdown_to_docx(md_path, docx_path, is_internal=False):
    doc = Document()
    
    # Portada básica
    titulo = "Plan de Seguridad Informática Integral"
    if is_internal:
        titulo += " (INTERNO - CON TRAZABILIDAD)"
    else:
        titulo += " (CLIENTE - EJECUTIVO)"
        
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(24)
    
    doc.add_paragraph("\nVersión: 0.1")
    doc.add_paragraph("Clasificación: Confidencial")
    doc.add_page_break()
    
    # TOC Placeholder
    doc.add_heading("Índice", level=1)
    doc.add_paragraph("[El Índice será generado por Microsoft Word. Haga clic derecho aquí y seleccione 'Actualizar campos']\n")
    doc.add_page_break()

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        
        if not line:
            idx += 1
            continue
            
        # Headers
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            text = sanitize_xml(line[level:].strip())
            # Ignorar el H1 principal porque ya hicimos portada
            if not (level == 1 and text == "Plan de Seguridad Informática Integral"):
                doc.add_heading(text, level=level)
            idx += 1
            if level <= 2:
                 doc.add_paragraph()
            continue
            
        # Listas
        if line.startswith('- ') or line.startswith('* '):
            text = sanitize_xml(line[2:].strip())
            doc.add_paragraph(text, style='List Bullet')
            idx += 1
            continue
            
        if line[0].isdigit() and len(line) > 1 and line[1] == '.':
            text = sanitize_xml(line[2:].strip())
            doc.add_paragraph(text, style='List Number')
            idx += 1
            continue
            
        # Tablas
        if line.startswith('|'):
            table_lines, new_idx = extract_table(lines, idx)
            if len(table_lines) > 2: # Tiene que tener cabecera, sep, y rows
                header = [c.strip() for c in table_lines[0].strip('|').split('|')]
                # idx 1 es el separador |---|---|
                data = []
                for row_line in table_lines[2:]:
                    cols = [c.strip() for c in row_line.strip('|').split('|')]
                    data.append(cols)
                    
                table = doc.add_table(rows=1, cols=len(header))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(header):
                    if i < len(hdr_cells):
                        hdr_cells[i].text = col_name.replace('**', '')
                        
                for row_data in data:
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row_data):
                        if i < len(row_cells):
                            # Limpiar un poco el formato MD
                            clean_val = sanitize_xml(val.replace('**', '').replace('`', ''))
                            row_cells[i].text = clean_val
                doc.add_paragraph()
            idx = new_idx
            continue
            
        # Párrafos normales
        if not line.startswith('---'):
            clean_line = sanitize_xml(line.replace('**', ''))
            doc.add_paragraph(clean_line)
            
        idx += 1

    # Si es el plan interno y existe el anexo, tratar de añadirlo (solo evidencia básica si Python-docx)
    annex_path = md_path.parent / 'ANNEX_evidence.md'
    if is_internal and annex_path.exists():
        doc.add_page_break()
        doc.add_heading("Anexo: Evidencias y Trazabilidad", level=1)
        with open(annex_path, 'r', encoding='utf-8') as f:
            for al in f:
                al = al.strip()
                if al.startswith('##'):
                    doc.add_heading(sanitize_xml(al.replace('#', '').strip()), level=2)
                elif al.startswith('- '):
                    doc.add_paragraph(sanitize_xml(al[2:]), style='List Bullet')
                elif al and not al.startswith('---'):
                    doc.add_paragraph(sanitize_xml(al.replace('**', '').replace('`', '')))

    doc.save(str(docx_path))
    print(f"  ✓ Exportado DOCX -> {docx_path.relative_to(PROJECT_ROOT)}")


def main():
    print("=" * 60)
    print("  EXPORTADOR DOCX (Fallback: python-docx)")
    print("=" * 60)
    
    cliente_md = PROJECT_ROOT / "deliverables" / "plan_seguridad_empresa_cliente" / "Plan_Seguridad_Empresa_CLIENTE.md"
    cliente_docx = PROJECT_ROOT / "deliverables" / "plan_seguridad_empresa_cliente" / "Plan_Seguridad_Empresa_CLIENTE.docx"
    
    interno_md = PROJECT_ROOT / "deliverables" / "plan_seguridad_empresa_interno" / "Plan_Seguridad_Empresa_INTERNO.md"
    interno_docx = PROJECT_ROOT / "deliverables" / "plan_seguridad_empresa_interno" / "Plan_Seguridad_Empresa_INTERNO.docx"
    
    didactico = PROJECT_ROOT / "exports" / "Plan_Seguridad_Informatica_Didactico_v1.md"
    didactico_docx = PROJECT_ROOT / "exports" / "Plan_Seguridad_Informatica_Didactico_v1.docx"
    
    if didactico.exists():
        parse_markdown_to_docx(didactico, didactico_docx, is_internal=False)
        os.remove(didactico)
        print("  ✓ Archivo Markdown Didactico eliminado, entregable único Creado.")
        
if __name__ == '__main__':
    main()
